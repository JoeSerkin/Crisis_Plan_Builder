"""Tests for DOCX export formatting."""

from __future__ import annotations

import pytest

pytest.importorskip("docx")

from cmp.render.docx_export import DocumentMetadata, markdown_to_docx


SAMPLE_MARKDOWN = """# Risk Register — Example Co

> **DRAFT — For consultant review.**

**Planning readiness score:** 80/100

## Tier 1 — Immediate Threat

| ID | Risk | Impact |
|----|------|--------|
| R-T1-001 | Industrial accident | critical |

- **Tier 1 Definition:** Immediate threat within 24 hours.
- Monitor emerging risks weekly.

1. Activate CMT.
2. Notify executive sponsor.
"""


def test_markdown_to_docx_creates_tables_and_lists(tmp_path) -> None:
    doc = markdown_to_docx(
        SAMPLE_MARKDOWN,
        metadata=DocumentMetadata(
            client_name="Example Co",
            engagement_id="example-mfg",
            document_title="Risk Register",
            industry="Manufacturing",
        ),
    )
    path = tmp_path / "risk_register.docx"
    doc.save(str(path))

    assert path.exists()
    assert path.stat().st_size > 3000

    from docx import Document

    loaded = Document(str(path))
    assert any("Risk Register" in paragraph.text for paragraph in loaded.paragraphs)
    assert loaded.tables
    assert loaded.tables[0].rows[0].cells[0].text == "ID"


def test_markdown_to_docx_skips_duplicate_title() -> None:
    doc = markdown_to_docx(
        "# Crisis Plan\n\nBody paragraph.",
        metadata=DocumentMetadata(document_title="Crisis Plan"),
    )
    titles = [p.text for p in doc.paragraphs if p.text == "Crisis Plan"]
    assert len(titles) == 1
    assert any(p.text == "Body paragraph." for p in doc.paragraphs)
