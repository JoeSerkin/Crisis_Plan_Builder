"""Extract text from uploaded documents and propose intake field updates."""

from __future__ import annotations

import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from cmp.intake.document_structured_extract import extract_structured_fields
from cmp.models.schemas import ClientIntake, RequirementGap

EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(
    r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{2,4}\)?[\s.-]?)?\d{3,4}[\s.-]?\d{3,4}"
)
URL_RE = re.compile(r"https?://[^\s)>]+", re.I)
EMPLOYEE_COUNT_RE = re.compile(
    r"(?:approximately|about|approx\.?|~)?\s*(\d[\d,]{1,6})\s+(?:employees|staff|headcount|FTEs?)\b",
    re.I,
)
YES_NO_RE = re.compile(r"\b(yes|no|not applicable|n/a)\b", re.I)
LABEL_VALUE_RE = re.compile(
    r"^[\s>*\-•]*(?P<label>.{3,80}?)\s*[:\-–—]\s+(?P<value>.+?)\s*$"
)
STOPWORDS = {
    "a",
    "an",
    "the",
    "and",
    "or",
    "of",
    "for",
    "to",
    "in",
    "on",
    "at",
    "by",
    "with",
    "is",
    "are",
    "be",
    "any",
    "all",
    "primary",
    "organization",
    "organizational",
    "please",
    "provide",
    "list",
    "describe",
    "briefly",
}


@dataclass
class ExtractionProposal:
    requirement_id: str
    field_path: str
    label: str
    proposed_value: Any
    confidence: str
    source_snippet: str
    source: str = "document"

    def to_dict(self) -> dict[str, Any]:
        return {
            "requirement_id": self.requirement_id,
            "field_path": self.field_path,
            "label": self.label,
            "proposed_value": self.proposed_value,
            "confidence": self.confidence,
            "source_snippet": self.source_snippet,
            "source": self.source,
        }


def extract_text_from_path(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".csv"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".json":
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        return _extract_pdf_text(path)
    if suffix == ".docx":
        return _extract_docx_text(path)
    raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")


def _extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError(
            "pypdf is required for PDF extraction. Install with: pip install -r requirements-v2.txt"
        ) from exc

    reader = PdfReader(str(path))
    chunks = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            chunks.append(text)
    return "\n\n".join(chunks)


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX extraction. Install with: pip install -r requirements-v2.txt"
        ) from exc

    document = Document(str(path))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _keywords(label: str) -> list[str]:
    tokens = re.findall(r"[A-Za-z0-9]+", label.lower())
    return [token for token in tokens if len(token) > 2 and token not in STOPWORDS]


def _line_matches_label(line: str, label: str) -> bool:
    lowered = line.lower()
    keys = _keywords(label)
    if not keys:
        return False
    hits = sum(1 for key in keys if key in lowered)
    return hits >= max(1, min(2, len(keys)))


def _normalize_scalar(value: str, field_path: str) -> Any:
    cleaned = value.strip().strip("•*-")
    if not cleaned:
        return cleaned

    lower_path = field_path.lower()
    if "email" in lower_path:
        match = EMAIL_RE.search(cleaned)
        return match.group(0) if match else cleaned
    if any(token in lower_path for token in ("phone", "hotline", "mobile", "fax")):
        match = PHONE_RE.search(cleaned)
        return match.group(0) if match else cleaned
    if field_path == "employees":
        match = EMPLOYEE_COUNT_RE.search(cleaned) or re.search(r"\b(\d[\d,]{1,6})\b", cleaned)
        if match:
            return int(match.group(1).replace(",", ""))
    if cleaned.lower() in {"yes", "no", "n/a", "not applicable"}:
        yn = YES_NO_RE.search(cleaned)
        if yn:
            token = yn.group(1).lower()
            if token in {"n/a", "not applicable"}:
                return "N/A"
            return token.capitalize()
    if field_path == "countries":
        parts = re.split(r",|/| and ", cleaned)
        countries = [part.strip(" .") for part in parts if part.strip(" .")]
        return countries if len(countries) > 1 else cleaned
    return cleaned


WEAK_VALUE_PATTERNS = (
    re.compile(r"^drills?$", re.I),
    re.compile(r"^emergency contingency and response plan$", re.I),
    re.compile(r"^risk profile and management chart$", re.I),
    re.compile(r"^make the plans\b", re.I),
    re.compile(r"^level one emergency\b", re.I),
)


def _is_quality_proposal(proposal: ExtractionProposal, text: str) -> bool:
    value = proposal.proposed_value
    if proposal.confidence == "high":
        return True
    if isinstance(value, str):
        cleaned = value.strip()
        if len(cleaned) < 4:
            return False
        if any(pattern.search(cleaned) for pattern in WEAK_VALUE_PATTERNS):
            return False
        if cleaned.upper() == cleaned and len(cleaned) < 80 and proposal.confidence == "low":
            return False
        first_line = next((line.strip() for line in text.splitlines() if line.strip()), "")
        if cleaned == first_line and proposal.field_path not in {
            "existing_crisis_plan",
            "pandemic_health_plan",
            "medical_emergency_response",
        }:
            return False
    return proposal.confidence != "low" or (
        isinstance(value, str) and len(value.strip()) >= 24
    )


def _extract_from_label_value_lines(text: str, gap: RequirementGap) -> ExtractionProposal | None:
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        match = LABEL_VALUE_RE.match(line)
        if not match:
            continue
        label = match.group("label")
        if not _line_matches_label(label, gap.label):
            continue
        value = _normalize_scalar(match.group("value"), gap.field_path)
        if isinstance(value, str) and not value:
            continue
        return ExtractionProposal(
            requirement_id=gap.requirement_id,
            field_path=gap.field_path,
            label=gap.label,
            proposed_value=value,
            confidence="medium",
            source_snippet=line[:240],
        )
    return None


def _extract_from_keyword_line(text: str, gap: RequirementGap) -> ExtractionProposal | None:
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line or not _line_matches_label(line, gap.label):
            continue
        if ":" in line or " - " in line:
            continue
        if len(line) > 240:
            continue
        value = _normalize_scalar(line, gap.field_path)
        if isinstance(value, str) and len(value) < 3:
            continue
        return ExtractionProposal(
            requirement_id=gap.requirement_id,
            field_path=gap.field_path,
            label=gap.label,
            proposed_value=value,
            confidence="low",
            source_snippet=line[:240],
        )
    return None


def _extract_pattern_value(text: str, gap: RequirementGap) -> ExtractionProposal | None:
    lower_path = gap.field_path.lower()
    if gap.field_path == "employees":
        match = EMPLOYEE_COUNT_RE.search(text)
        if match:
            return ExtractionProposal(
                requirement_id=gap.requirement_id,
                field_path=gap.field_path,
                label=gap.label,
                proposed_value=int(match.group(1).replace(",", "")),
                confidence="medium",
                source_snippet=match.group(0)[:240],
            )
    if "email" in lower_path:
        match = EMAIL_RE.search(text)
        if match:
            return ExtractionProposal(
                requirement_id=gap.requirement_id,
                field_path=gap.field_path,
                label=gap.label,
                proposed_value=match.group(0),
                confidence="medium",
                source_snippet=match.group(0),
            )
    if any(token in lower_path for token in ("phone", "hotline", "mobile")):
        match = PHONE_RE.search(text)
        if match:
            return ExtractionProposal(
                requirement_id=gap.requirement_id,
                field_path=gap.field_path,
                label=gap.label,
                proposed_value=match.group(0),
                confidence="low",
                source_snippet=match.group(0),
            )
    return None


def propose_updates_from_text(
    text: str,
    gaps: list[RequirementGap],
    *,
    intake: ClientIntake | None = None,
) -> list[ExtractionProposal]:
    gap_by_field = {gap.field_path: gap for gap in gaps}
    proposals: list[ExtractionProposal] = []
    seen_fields: set[str] = set()

    for field_path, value, confidence, snippet in extract_structured_fields(text):
        gap = gap_by_field.get(field_path)
        if gap is None:
            continue
        if field_path in seen_fields:
            continue
        proposals.append(
            ExtractionProposal(
                requirement_id=gap.requirement_id,
                field_path=field_path,
                label=gap.label,
                proposed_value=value,
                confidence=confidence,
                source_snippet=snippet,
                source="structured",
            )
        )
        seen_fields.add(field_path)

    for gap in gaps:
        if gap.field_path in seen_fields:
            continue
        if intake and intake.is_field_present(gap.field_path):
            continue
        proposal = (
            _extract_from_label_value_lines(text, gap)
            or _extract_from_keyword_line(text, gap)
            or _extract_pattern_value(text, gap)
        )
        if proposal and _is_quality_proposal(proposal, text):
            proposals.append(proposal)
            seen_fields.add(proposal.field_path)

    return proposals


def propose_updates_from_json_text(text: str, gaps: list[RequirementGap]) -> list[ExtractionProposal]:
    try:
        payload = json.loads(text)
    except json.JSONDecodeError:
        return []

    if not isinstance(payload, dict):
        return []

    gap_by_field = {gap.field_path: gap for gap in gaps}
    proposals: list[ExtractionProposal] = []

    def walk(prefix: str, value: Any) -> None:
        if isinstance(value, dict):
            for key, nested in value.items():
                path = f"{prefix}.{key}" if prefix else key
                walk(path, nested)
            return
        gap = gap_by_field.get(prefix)
        if gap is None:
            return
        if value in (None, "", [], {}):
            return
        proposals.append(
            ExtractionProposal(
                requirement_id=gap.requirement_id,
                field_path=gap.field_path,
                label=gap.label,
                proposed_value=value,
                confidence="high",
                source_snippet=f"JSON field {gap.field_path}",
                source="json_upload",
            )
        )

    for key, value in payload.items():
        if key == "additional_context" and isinstance(value, dict):
            for sub_key, sub_value in value.items():
                walk(sub_key, sub_value)
        else:
            walk(key, value)

    deduped: dict[str, ExtractionProposal] = {}
    for proposal in proposals:
        deduped[proposal.field_path] = proposal
    return list(deduped.values())


def propose_updates_from_document(
    path: Path,
    gaps: list[RequirementGap],
    *,
    intake: ClientIntake | None = None,
) -> tuple[str, list[ExtractionProposal]]:
    text = extract_text_from_path(path)
    if path.suffix.lower() == ".json":
        proposals = propose_updates_from_json_text(text, gaps)
    else:
        proposals = propose_updates_from_text(text, gaps, intake=intake)
    return text, proposals
