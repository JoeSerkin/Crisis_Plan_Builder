"""Convert markdown deliverables to formatted DOCX for client delivery."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING, Any, Iterator

from cmp.models.requirements import repo_root
from cmp.storage.engagement_store import EngagementStore

if TYPE_CHECKING:
    from docx.document import Document
    from docx.text.paragraph import Paragraph

_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^(\s*)[-*]\s+(.*)$")
_NUMBERED = re.compile(r"^(\s*)\d+\.\s+(.*)$")
_TABLE_ROW = re.compile(r"^\|(.+)\|$")
_SEPARATOR_CELL = re.compile(r"^:?-+:?$")
_BOLD = re.compile(r"\*\*(.+?)\*\*")
_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")

BRAND_NAVY = "1E3A5F"
BRAND_SLATE = "5A6570"
CALLOUT_FILL = "EEF2F7"
TABLE_HEADER_FILL = "1E3A5F"
TABLE_ALT_FILL = "F4F6F8"


def _rgb(hex_color: str):
    _, _, _, _, _, _, RGBColor = _require_docx()
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


@dataclass
class DocumentMetadata:
    client_name: str = ""
    engagement_id: str = ""
    document_title: str = ""
    industry: str = ""


def _require_docx():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX export. Install with: pip install -r requirements-v2.txt"
        ) from exc
    return Document, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Inches, Pt, RGBColor


def _shade_paragraph(paragraph: Paragraph, fill: str) -> None:
    _, _, OxmlElement, qn, *_ = _require_docx()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    paragraph._p.get_or_add_pPr().append(shading)


def _set_cell_shading(cell, fill: str) -> None:
    _, _, OxmlElement, qn, *_ = _require_docx()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _add_formatted_runs(paragraph: Paragraph, text: str, *, bold_default: bool = False) -> None:
    _, _, _, _, _, Pt, RGBColor = _require_docx()
    if not text:
        return

    parts: list[tuple[str, bool, bool]] = []
    cursor = 0
    for match in re.finditer(r"(\*\*.+?\*\*|\*.+?\*)", text):
        if match.start() > cursor:
            parts.append((text[cursor : match.start()], bold_default, False))
        token = match.group(0)
        if token.startswith("**"):
            parts.append((token[2:-2], True, False))
        else:
            parts.append((token[1:-1], bold_default, True))
        cursor = match.end()
    if cursor < len(text):
        parts.append((text[cursor:], bold_default, False))

    if not parts:
        parts = [(text, bold_default, False)]

    for content, bold, italic in parts:
        run = paragraph.add_run(content)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(11)


def _add_paragraph(
    doc: Document,
    text: str,
    *,
    style: str | None = None,
    callout: bool = False,
) -> Paragraph:
    paragraph = doc.add_paragraph(style=style)
    _add_formatted_runs(paragraph, text)
    if callout:
        _shade_paragraph(paragraph, CALLOUT_FILL)
    return paragraph


def _parse_table_cells(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_separator_row(cells: list[str]) -> bool:
    return bool(cells) and all(_SEPARATOR_CELL.match(cell) for cell in cells)


def _iter_markdown_blocks(lines: list[str]) -> Iterator[tuple[str, Any]]:
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        heading = _HEADING.match(line)
        if heading:
            yield ("heading", (len(heading.group(1)), heading.group(2).strip()))
            index += 1
            continue

        if stripped in {"---", "***", "___"}:
            yield ("hr", None)
            index += 1
            continue

        if stripped.startswith(">"):
            block_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                block_lines.append(lines[index].strip()[1:].strip())
                index += 1
            yield ("blockquote", " ".join(block_lines))
            continue

        if _TABLE_ROW.match(line):
            rows: list[list[str]] = []
            while index < len(lines) and _TABLE_ROW.match(lines[index]):
                cells = _parse_table_cells(lines[index])
                if not _is_separator_row(cells):
                    rows.append(cells)
                index += 1
            if rows:
                yield ("table", rows)
            continue

        if _BULLET.match(line):
            items: list[tuple[int, str]] = []
            while index < len(lines):
                bullet = _BULLET.match(lines[index])
                if not bullet:
                    break
                indent = len(bullet.group(1).replace("\t", "    ")) // 2
                items.append((indent, bullet.group(2).strip()))
                index += 1
            yield ("bullet_list", items)
            continue

        numbered = _NUMBERED.match(line)
        if numbered:
            items = []
            while index < len(lines):
                item = _NUMBERED.match(lines[index])
                if not item:
                    break
                items.append(item.group(2).strip())
                index += 1
            yield ("numbered_list", items)
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            nxt = lines[index].strip()
            if (
                not nxt
                or _HEADING.match(lines[index])
                or nxt.startswith(">")
                or _TABLE_ROW.match(lines[index])
                or _BULLET.match(lines[index])
                or _NUMBERED.match(lines[index])
                or nxt in {"---", "***", "___"}
            ):
                break
            paragraph_lines.append(nxt)
            index += 1
        yield ("paragraph", " ".join(paragraph_lines))


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    _, _, _, _, _, Pt, RGBColor = _require_docx()
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            cell = table.rows[row_index].cells[col_index]
            value = row[col_index] if col_index < len(row) else ""
            cell.text = ""
            paragraph = cell.paragraphs[0]
            _add_formatted_runs(paragraph, value, bold_default=row_index == 0)
            for run in paragraph.runs:
                run.font.size = Pt(10)
                if row_index == 0:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            if row_index == 0:
                _set_cell_shading(cell, TABLE_HEADER_FILL)
            elif row_index % 2 == 0:
                _set_cell_shading(cell, TABLE_ALT_FILL)

    doc.add_paragraph("")


def _configure_document(doc: Document, metadata: DocumentMetadata) -> None:
    Document, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Inches, Pt, RGBColor = _require_docx()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

    header = section.header.paragraphs[0]
    header.text = ""
    header_run = header.add_run(metadata.document_title or "Crisis Management Deliverable")
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = _rgb(BRAND_SLATE)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    left = footer.add_run("DRAFT — Consultant review required")
    left.font.size = Pt(9)
    left.font.color.rgb = _rgb(BRAND_SLATE)
    footer.add_run("    ")
    page_field = footer.add_run()
    page_field.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    page_field._r.append(fld)


def _add_cover_block(doc: Document, metadata: DocumentMetadata) -> None:
    _, WD_ALIGN_PARAGRAPH, _, _, _, Pt, RGBColor = _require_docx()
    if metadata.client_name:
        client = doc.add_paragraph()
        client.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = client.add_run(metadata.client_name)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = _rgb(BRAND_NAVY)

    if metadata.document_title:
        title = doc.add_heading(metadata.document_title, level=0)
        for run in title.runs:
            run.font.color.rgb = _rgb(BRAND_NAVY)

    meta_bits = [bit for bit in (metadata.industry, metadata.engagement_id) if bit]
    if meta_bits:
        subtitle = doc.add_paragraph(" · ".join(meta_bits))
        subtitle.runs[0].font.size = Pt(10)
        subtitle.runs[0].font.color.rgb = _rgb(BRAND_SLATE)

    doc.add_paragraph("")


def markdown_to_docx(
    markdown_text: str,
    *,
    metadata: DocumentMetadata | None = None,
    include_cover: bool = True,
) -> Document:
    Document, _, _, _, _, Pt, RGBColor = _require_docx()
    doc = Document()
    meta = metadata or DocumentMetadata()
    _configure_document(doc, meta)

    if include_cover and meta.document_title:
        _add_cover_block(doc, meta)

    skip_leading_title = include_cover and markdown_text.lstrip().startswith("# ")
    skipped_title = False

    for block_type, payload in _iter_markdown_blocks(markdown_text.splitlines()):
        if block_type == "heading":
            level, text = payload
            if skip_leading_title and not skipped_title and level == 1:
                skipped_title = True
                continue
            doc.add_heading(text, level=min(level, 4))
            continue

        if block_type == "hr":
            doc.add_paragraph("")
            continue

        if block_type == "blockquote":
            _add_paragraph(doc, str(payload), callout=True)
            continue

        if block_type == "table":
            _add_table(doc, payload)
            continue

        if block_type == "bullet_list":
            for indent, text in payload:
                try:
                    style = "List Bullet" if indent == 0 else "List Bullet 2"
                    paragraph = doc.add_paragraph(style=style)
                except KeyError:
                    paragraph = doc.add_paragraph(style="List Bullet")
                _add_formatted_runs(paragraph, text)
            continue

        if block_type == "numbered_list":
            for text in payload:
                paragraph = doc.add_paragraph(style="List Number")
                _add_formatted_runs(paragraph, text)
            continue

        if block_type == "paragraph":
            _add_paragraph(doc, str(payload))

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.size is None:
                run.font.size = Pt(11)

    return doc


def export_engagement_docx(
    engagement_id: str,
    output_dir: Path | None = None,
    store: EngagementStore | None = None,
) -> dict[str, Path]:
    source_dir = repo_root() / "output" / engagement_id
    if not source_dir.exists():
        raise FileNotFoundError(f"No deliverables found for engagement {engagement_id}")

    engagement_store = store or EngagementStore()
    record = engagement_store.get_engagement(engagement_id)
    intake = engagement_store.load_intake(engagement_id)

    target_dir = output_dir or (source_dir / "docx")
    target_dir.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}

    markdown_files = [
        path for path in source_dir.rglob("*.md") if "docx" not in path.parts
    ]
    if not markdown_files:
        raise FileNotFoundError(f"No markdown deliverables in {source_dir}")

    for md_path in markdown_files:
        rel = md_path.relative_to(source_dir)
        docx_path = target_dir / rel.with_suffix(".docx")
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        content = md_path.read_text(encoding="utf-8")
        title = md_path.stem.replace("_", " ").title()
        metadata = DocumentMetadata(
            client_name=intake.company_name if intake else (record.client_name if record else ""),
            engagement_id=engagement_id,
            document_title=title,
            industry=(record.industry if record else intake.industry if intake else ""),
        )
        doc = markdown_to_docx(content, metadata=metadata)
        doc.save(str(docx_path))
        paths[str(rel.with_suffix(".docx")).replace("\\", "/")] = docx_path

    return paths
